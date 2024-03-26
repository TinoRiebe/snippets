from docx import Document

doc = 'MOG2.docx'


def extract_table_of_contents(doc_path):
    try:
        toc = []
        document = Document(doc_path)

        for paragraph in document.paragraphs:
            print(paragraph.style.name)
            if paragraph.style.name.startswith('Heading'):
                text = paragraph.text.strip()
                toc.append(paragraph.text)
        return toc

    except Exception as e:
        print('Fehler: ', e)
        return None

def extract_kapitals_and_chapters(doc):
    try:
        document = Document(doc)
        kapitel_und_splitts = {}
        act_kap = None
        act_content = []
        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph.style.name.startswith('Heading') and paragraph_text:
                if act_kap is not None:
                    kapitel_und_splitts[act_kap] = dict(enumerate(act_content))
                act_kap = paragraph_text
                act_content = []
            else:
                act_content.append(paragraph.text)
        if act_kap is not None:
            kapitel_und_splitts[act_kap] = dict(enumerate(act_content))
        return kapitel_und_splitts

    except Exception as e:
        print('Fehler: ', e)
        return None


toc = extract_table_of_contents(doc)
print(toc)
content = extract_kapitals_and_chapters(doc)
print(content)
